import streamlit as st
import fitz  # pymupdf
import re
import pandas as pd
from docx import Document
from num2words import num2words
from datetime import datetime, timedelta
import calendar
import io
import base64
import os
from resources import CODIGOS_BRUTO, CODIGOS_DEDUCCIONES, MOTIVOS, TOPE_MAXIMO_PRESTAMO

# Configuración de la página
st.set_page_config(
    page_title="Sistema de Adelantos Haberes",
    page_icon="💰",
    layout="wide"
)

# Estilos CSS personalizados
st.markdown("""
    <style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .stDataFrame {
        width: 100%;
    }
    .stButton>button {
        width: 100%;
    }
    </style>
""", unsafe_allow_html=True)

# Funciones auxiliares (mantenidas del código original)
def extraer_sueldos(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()

        lines = text.splitlines()
        
        # Extraer nombre (buscando después de "Apellido y Nombre:")
        nombre = None
        for i, line in enumerate(lines):
            if "Apellido y Nombre:" in line:
                # Buscar la siguiente línea que contenga una coma (formato "Apellido, Nombre")
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if "," in next_line and not any(field in next_line for field in ["Categoria:", "Cargo:", "Egreso:", "Codigo", "Concepto"]):
                        nombre = next_line
                        # Convertir de "Apellido, Nombre" a "Nombre Apellido"
                        if "," in nombre:
                            apellido, nombre_persona = nombre.split(",", 1)
                            nombre = f"{nombre_persona.strip()} {apellido.strip()}"
                        break
                    j += 1
                break
        
        # Extraer montos
        monto_regex = re.compile(r'^\s*(\d{1,3}(?:\.\d{3})*,\d{2})\s*$')
        valores = [match.group(1) for line in lines if (match := monto_regex.match(line))]

        if len(valores) < 2:
            st.error("No se encontraron suficientes montos claros para bruto/neto")
            return None, None, None

        valores_f = [float(v.replace('.', '').replace(',', '.')) for v in valores]
        sueldo_neto = valores_f[-1]
        candidatos = [v for v in valores_f[-6:] if v > 1_000_000]
        if not candidatos:
            st.error("No se detectó un valor alto para el sueldo bruto.")
            return None, None, None
        sueldo_bruto = max(candidatos)

        return sueldo_bruto, sueldo_neto, nombre

    except Exception as e:
        st.error(f"Error al procesar PDF: {e}")
        return None, None, None

def calcular_bloques_forzado(pdf_path):
    from resources import CODIGOS_BRUTO, CODIGOS_DEDUCCIONES
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()

        lines = text.splitlines()
        bruto = 0.0
        deducciones = 0.0
        detectados = []

        def es_monto(s):
            return re.match(r'^-?\d{1,3}(?:\.\d{3})*,\d{2}$', s)
        def es_cantidad(s):
            return re.match(r'^\d{1,3}(?:\.\d{3})*,\d{2}$', s)

        # Buscar la sección de conceptos
        inicio_conceptos = False
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            # Detectar inicio de la sección de conceptos
            if line == "Codigo":
                inicio_conceptos = True
                i += 1
                continue
            if inicio_conceptos:
                for codigo in CODIGOS_BRUTO.keys():
                    # Solo considerar líneas que empiezan por el código, espacio y una letra (no número ni coma)
                    if re.match(rf'^{codigo} [A-Za-z]', line):
                        # Caso 1: cantidad y luego monto
                        if i + 2 < len(lines) and es_cantidad(lines[i+1].strip()) and es_monto(lines[i+2].strip()):
                            valor_str = lines[i+2].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            bruto += valor
                            detectados.append((codigo, valor, "REM", line))
                        # Caso 2: monto directo
                        elif i + 1 < len(lines) and es_monto(lines[i+1].strip()):
                            valor_str = lines[i+1].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            bruto += valor
                            detectados.append((codigo, valor, "REM", line))
                        break
                for codigo in CODIGOS_DEDUCCIONES.keys():
                    if re.match(rf'^{codigo} [A-Za-z]', line):
                        if i + 2 < len(lines) and es_cantidad(lines[i+1].strip()) and es_monto(lines[i+2].strip()):
                            valor_str = lines[i+2].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            deducciones += valor
                            detectados.append((codigo, valor, "DED", line))
                        elif i + 1 < len(lines) and es_monto(lines[i+1].strip()):
                            valor_str = lines[i+1].strip()
                            valor = float(valor_str.replace('.', '').replace(',', '.'))
                            deducciones += valor
                            detectados.append((codigo, valor, "DED", line))
                        break
            i += 1

        neto = bruto - deducciones
        return round(bruto, 2), round(deducciones, 2), round(neto, 2), detectados

    except Exception as e:
        st.error(f"Error al procesar PDF: {e}")
        return None, None, None, None

def calcular_cuota(monto, cuotas, tasa_anual):
    tasa_mensual = (tasa_anual / 100) / 12
    if tasa_mensual == 0:
        return monto / cuotas
    cuota = monto * (tasa_mensual * (1 + tasa_mensual)**cuotas) / ((1 + tasa_mensual)**cuotas - 1)
    return cuota

def generar_cuadro_amortizacion(monto, cuotas, tasa_anual):
    tasa_mensual = (tasa_anual / 100) / 12
    cuota_total = calcular_cuota(monto, cuotas, tasa_anual)
    saldo = monto
    cuadro = []
    for i in range(1, cuotas + 1):
        interes = saldo * tasa_mensual
        amortizacion = cuota_total - interes
        saldo -= amortizacion
        cuadro.append({
            "Cuota N°": i,
            "Cuota total ($)": round(cuota_total, 2),
            "Interés ($)": round(interes, 2),
            "Amortización ($)": round(amortizacion, 2),
            "Saldo restante ($)": round(saldo if saldo > 0 else 0, 2)
        })
    return pd.DataFrame(cuadro)

def generar_nota(monto, cuotas, tasa_final, cuota, fecha, nombre, area, sector, motivo, motivo_detallado, puesto, neto):
    def formatear_fecha_larga(fecha):
        meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                 'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
        return f"{fecha.day} de {meses[fecha.month - 1]} del {fecha.year}"

    def tercer_viernes(fecha_base):
        year = fecha_base.year
        month = fecha_base.month
        count = 0
        for day in range(1, 32):
            try:
                fecha = datetime(year, month, day)
                if fecha.weekday() == 4:
                    count += 1
                    if count == 3:
                        return fecha
            except ValueError:
                break
        return fecha_base

    def ultimo_dia_habil_del_mes(fecha_base):
        anio, mes = fecha_base.year, fecha_base.month
        ultimo_dia = calendar.monthrange(anio, mes)[1]
        venc = datetime(anio, mes, ultimo_dia)
        while venc.weekday() >= 5:
            venc -= timedelta(days=1)
        return venc

    try:
        fecha_directorio = tercer_viernes(fecha)
        vencimiento = ultimo_dia_habil_del_mes(fecha)
        texto_letras = num2words(monto, lang='es').replace("uno", "un").capitalize() + " pesos"
        neto_menos_cuota = neto - cuota

        datos = {
            "<nombre>": nombre,
            "<area>": area,
            "<sector>": sector,
            "<fecha>": formatear_fecha_larga(fecha),
            "<fecha_directorio>": formatear_fecha_larga(fecha_directorio),
            "<monto>": f"${monto:,.2f}",
            "<cuotas>": str(cuotas),
            "<motivo>": motivo,
            "<motivo_detallado>": motivo_detallado,
            "<monto_en_letras>": texto_letras,
            "<tasa>": f"{tasa_final:.2f}%",
            "<vencimiento>": formatear_fecha_larga(vencimiento),
            "<puesto>": puesto,
            "<neto_menos_cuota>": f"${neto_menos_cuota:,.2f}"
        }

        # Buscar la plantilla
        plantilla = None
        for archivo in os.listdir(os.getcwd()):
            if archivo.endswith(".docx") and "nota" in archivo.lower():
                doc_test = Document(archivo)
                texts = [p.text for p in doc_test.paragraphs]
                texts += [c.text for t in doc_test.tables for r in t.rows for c in r.cells]
                if any("<" in t and ">" in t for t in texts):
                    plantilla = archivo
                    break

        if not plantilla:
            st.error("❌ No se encontró una plantilla con '<>' en la carpeta.")
            return None

        doc = Document(plantilla)

        # Reemplazar marcadores en párrafos
        for p in doc.paragraphs:
            for k, v in datos.items():
                if k in p.text:
                    for r in p.runs:
                        r.text = r.text.replace(k, v)
                # Reemplazo global por si el tag está partido en runs
                if k in p.text:
                    p.text = p.text.replace(k, v)

        # Reemplazar marcadores en tablas
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for k, v in datos.items():
                            if k in p.text:
                                for run in p.runs:
                                    run.text = run.text.replace(k, v)
                            # Reemplazo global por si el tag está partido en runs
                            if k in p.text:
                                p.text = p.text.replace(k, v)

        # Reemplazar marcadores en texto plano de celdas (por si acaso)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for k, v in datos.items():
                        if k in c.text:
                            c.text = c.text.replace(k, v)

        # Intentar agregar la tabla de amortización
        try:
            df_amort = generar_cuadro_amortizacion(monto, cuotas, tasa_final)
            for i, p in enumerate(doc.paragraphs):
                if "<cuadro_amortizacion>" in p.text:
                    p.text = p.text.replace("<cuadro_amortizacion>", "")
                    table = doc.add_table(rows=1, cols=len(df_amort.columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for j, col in enumerate(df_amort.columns):
                        hdr_cells[j].text = str(col)
                    for _, row in df_amort.iterrows():
                        row_cells = table.add_row().cells
                        for j, val in enumerate(row):
                            row_cells[j].text = str(val)
                    p._p.addnext(table._tbl)
                    break
        except Exception as e:
            st.error(f"No se pudo insertar la tabla: {e}")

        # Guardar en memoria
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes

    except Exception as e:
        st.error(f"❌ Error al generar nota: {e}")
        return None

# Interfaz principal
st.title("Sistema de Adelantos Haberes")

# Instrucciones justo después del título
st.markdown("""
### Instrucciones
1. Cargue su recibo de sueldo en formato PDF
2. Complete sus datos personales en el panel lateral
3. Ingrese los datos del préstamo deseado
4. Simule el préstamo para ver el cuadro de amortización
5. Genere y descargue la nota de solicitud
""")

# Sidebar para datos del usuario
with st.sidebar:
    st.header("Datos del Usuario")
    
    # Inicializar el estado del nombre si no existe
    if 'nombre_usuario' not in st.session_state:
        st.session_state.nombre_usuario = ""
    
    # Si hay un nombre en los parámetros de consulta, usarlo
    if 'nombre' in st.query_params:
        st.session_state.nombre_usuario = st.query_params['nombre']
    
    nombre = st.text_input("Nombre completo", value=st.session_state.nombre_usuario)
    area = st.text_input("Área")
    sector = st.text_input("Sector")
    motivo = st.selectbox("Motivo", MOTIVOS)
    motivo_detallado = st.text_area("Motivo de la solicitud")
    puesto = st.text_input("Puesto")

# Sección de entrada de datos en columnas
st.header("Datos del Préstamo")
col1, col2 = st.columns(2)

with col1:
    st.subheader("Carga de Recibo de Sueldo")
    uploaded_file = st.file_uploader("Seleccione el recibo de sueldo (PDF)", type=['pdf'])
    
    if uploaded_file is not None:
        # Guardar el archivo temporalmente
        with open("temp.pdf", "wb") as f:
            f.write(uploaded_file.getvalue())
        
        # Extraer desglose y cálculos correctos
        bruto, deducciones, neto, detectados = calcular_bloques_forzado("temp.pdf")
        
        # Extraer nombre (opcional, solo para el campo)
        _, _, nombre_detectado = extraer_sueldos("temp.pdf")
        
        if bruto is not None and neto is not None:
            st.session_state['bruto'] = bruto
            st.session_state['neto'] = neto
            
            # Mostrar solo los totales
            st.success(f"Sueldo bruto: ${bruto:,.2f}")
            st.success(f"Sueldo neto: ${neto:,.2f}")
            
            # Si se detectó un nombre, actualizar el campo inmediatamente
            if nombre_detectado and nombre_detectado != st.session_state.nombre_usuario:
                st.session_state.nombre_usuario = nombre_detectado
                st.query_params["nombre"] = nombre_detectado
                st.rerun()  # Forzar la actualización de la interfaz
        else:
            st.error("No se pudieron extraer los datos del PDF. Por favor, ingréselos manualmente.")
            bruto = st.number_input("Sueldo bruto", min_value=0.0, step=1000.0)
            neto = st.number_input("Sueldo neto", min_value=0.0, step=1000.0)
            st.session_state['bruto'] = bruto
            st.session_state['neto'] = neto

with col2:
    st.subheader("Parámetros del Préstamo")
    
    # Inicializar el estado de la simulación
    if 'simulacion_realizada' not in st.session_state:
        st.session_state.simulacion_realizada = False
    if 'nota_generada' not in st.session_state:
        st.session_state.nota_generada = False
    
    monto = st.number_input(
        "Monto solicitado ($)",
        min_value=0.0,
        step=1000.0,
        format="%g",
        help="Ingrese el monto en pesos"
    )
    
    # Formatear el monto para mostrar
    monto_formateado = f"${monto:,.2f}"
    st.markdown(f"**Monto ingresado:** {monto_formateado}")
    
    cuotas = st.number_input("Cantidad de cuotas", min_value=1, max_value=18, step=1)
    tasa_anual = st.number_input("Tasa anual (%)", min_value=0.0, step=0.1)
    tasa_mensual = tasa_anual / 12
    fecha = st.date_input("Fecha de solicitud")

# Sección de simulación a ancho completo
st.header("Simulación")
if st.button("Simular", key="simular_button"):
    if 'bruto' not in st.session_state or 'neto' not in st.session_state:
        st.error("Por favor, cargue primero el recibo de sueldo")
        st.stop()

    bruto = st.session_state['bruto']
    neto = st.session_state['neto']

    # Validaciones
    if cuotas < 1 or cuotas > 18:
        st.error("La cantidad de cuotas debe ser entre 1 y 18.")
        st.stop()
    
    if monto > TOPE_MAXIMO_PRESTAMO:
        st.error(f"El monto excede el tope máximo permitido de ${TOPE_MAXIMO_PRESTAMO:,.2f}.")
        st.stop()
    
    if monto > 3 * bruto:
        st.error("El monto excede 3 veces el sueldo bruto.")
        st.stop()

    cuota = calcular_cuota(monto, cuotas, tasa_anual)

    if cuota > 0.3 * neto:
        st.error("La cuota mensual excede el 30% del sueldo neto.")
        st.stop()

    # Mostrar resumen
    st.subheader("Resumen de la simulación")
    col_resumen1, col_resumen2 = st.columns(2)
    with col_resumen1:
        st.write(f"Monto solicitado: ${monto:,.2f}")
        st.write(f"Cantidad de cuotas: {cuotas}")
        st.write(f"Cuota mensual estimada: ${cuota:,.2f}")
    with col_resumen2:
        st.write(f"Tasa anual: {tasa_anual:.2f}%")
        st.write(f"Tasa mensual: {tasa_mensual:.2f}%")

    # Generar cuadro de amortización
    st.subheader("Cuadro de Amortización")
    df_amort = generar_cuadro_amortizacion(monto, cuotas, tasa_anual)
    st.dataframe(
        df_amort,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Cuota N°": st.column_config.NumberColumn("Cuota N°", format="%d"),
            "Cuota total ($)": st.column_config.NumberColumn("Cuota total ($)", format="$%.2f"),
            "Interés ($)": st.column_config.NumberColumn("Interés ($)", format="$%.2f"),
            "Amortización ($)": st.column_config.NumberColumn("Amortización ($)", format="$%.2f"),
            "Saldo restante ($)": st.column_config.NumberColumn("Saldo restante ($)", format="$%.2f")
        }
    )

    # Guardar los datos de la simulación en el estado
    st.session_state.simulacion_realizada = True
    st.session_state.datos_simulacion = {
        'monto': monto,
        'cuotas': cuotas,
        'tasa_anual': tasa_anual,
        'cuota': cuota,
        'fecha': fecha,
        'neto': neto
    }

# Mostrar botón de generación de nota solo si la simulación fue exitosa
if st.session_state.simulacion_realizada:
    st.markdown("---")
    st.subheader("Generación de Nota")
    if st.button("Generar Nota", key="generar_nota_button"):
        if not all([nombre, area, sector, motivo, motivo_detallado, puesto]):
            st.error("Por favor complete todos los datos del usuario en el panel lateral")
            st.stop()

        datos = st.session_state.datos_simulacion
        docx_bytes = generar_nota(
            datos['monto'], datos['cuotas'], datos['tasa_anual'],
            datos['cuota'], datos['fecha'],
            nombre, area, sector, motivo, motivo_detallado, puesto, datos['neto']
        )

        if docx_bytes is not None:
            # Crear botón de descarga
            b64 = base64.b64encode(docx_bytes.getvalue()).decode()
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="nota.docx">Descargar Nota de Solicitud</a>'
            st.markdown(href, unsafe_allow_html=True)
            st.success("✅ Nota generada correctamente. Haga clic en el enlace para descargarla.")
            st.session_state.nota_generada = True
        else:
            st.error("❌ No se pudo generar la nota. Por favor, intente nuevamente.") 
